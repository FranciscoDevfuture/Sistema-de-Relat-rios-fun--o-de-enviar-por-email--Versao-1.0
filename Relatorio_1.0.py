#importando bibliotecas necessárias
import json
import smtplib
from datetime import datetime
from docx import Document
from email.message import EmailMessage

# Lista de relatórios
relatorios = []

# ---------- Funções Utilitárias ----------

def validar_data(data):
    """Valida o formato da data (DD/MM/AAAA)."""
    try:
        datetime.strptime(data, "%d/%m/%Y")
        return True
    except ValueError:
        return False
#Essas duas funções carregar_json() e salvar_json() 
# são responsáveis por manter os dados dos relatórios salvos 
# de forma permanente em um arquivo .json, permitindo que as informações
# sejam carregadas novamente quando o sistema for reaberto.
def carregar_json():
    """Carrega os relatórios salvos em JSON."""
    global relatorios
    try:
        with open('relatorios.json', 'r', encoding='utf-8') as f:
            relatorios.extend(json.load(f))
    except FileNotFoundError:
        pass

def salvar_json():
    """Salva os relatórios no arquivo JSON."""
    with open('relatorios.json', 'w', encoding='utf-8') as f:
        json.dump(relatorios, f, ensure_ascii=False, indent=4)

# ---------- Funcionalidades do Sistema ----------

def cadastrar_relatorio():
    """Cadastra um novo relatório."""
    relatorio = {}
    relatorio['id'] = relatorios[-1]['id'] + 1 if relatorios else 1
    relatorio['titulo'] = input("Digite o título do relatório: ")
    relatorio['descricao'] = input("Digite a descrição do relatório: ")

    while True:
        data = input("Digite a data do relatório (DD/MM/AAAA): ")
        if validar_data(data):
            relatorio['data'] = data
            break
        else:
            print("Data inválida. Tente novamente.")

    relatorios.append(relatorio)
    salvar_json()
    print("✅ Relatório cadastrado com sucesso!")

def listar_relatorios():
    """Lista todos os relatórios cadastrados."""
    if not relatorios:
        print("⚠️ Nenhum relatório cadastrado.")
        return

    print("\n📄 Relatórios cadastrados:")
    for r in relatorios:
        print(f"ID: {r['id']} | Título: {r['titulo']} | Data: {r['data']}")

def buscar_relatorio_por_id():
    """Busca um relatório pelo ID."""
    try:
        id_relatorio = int(input("Digite o ID do relatório: "))
    except ValueError:
        print("❌ ID inválido. Use apenas números.")
        return

    for r in relatorios:
        if r['id'] == id_relatorio:
            print(f"\n🔍 Relatório encontrado:")
            print(f"ID: {r['id']}\nTítulo: {r['titulo']}\nDescrição: {r['descricao']}\nData: {r['data']}")
            return

    print("❌ Relatório não encontrado.")

def buscar_por_titulo():
    """Busca relatórios contendo parte do título."""
    termo = input("Digite parte do título: ").lower()
    encontrados = [r for r in relatorios if termo in r['titulo'].lower()]

    if encontrados:
        print(f"\n🔍 {len(encontrados)} relatório(s) encontrado(s):")
        for r in encontrados:
            print(f"ID: {r['id']} | Título: {r['titulo']} | Data: {r['data']}")
    else:
        print("❌ Nenhum relatório encontrado com esse título.")

def salvar_em_docx():
    """Salva os relatórios em um arquivo Word."""
    if not relatorios:
        print("⚠️ Nenhum relatório para salvar.")
        return

    doc = Document()
    doc.add_heading('Relatórios de Ronda Cadastrados', level=1)

    for r in relatorios:
        doc.add_heading(r['titulo'], level=2)
        doc.add_paragraph(f"ID: {r['id']}")
        doc.add_paragraph(f"Descrição: {r['descricao']}")
        #doc.add_paragraph(f"Data: {r['data']}")
        doc.add_paragraph("-" * 40)

    doc.save('relatorios.docx')
    print("📁 Relatórios salvos em 'relatorios.docx'.")

def salvar_em_txt():
    """Salva os relatórios em um arquivo .txt."""
    if not relatorios:
        print("⚠️ Nenhum relatório para salvar.")
        return

    with open('relatorios.txt', 'w', encoding='utf-8') as f:
        for r in relatorios:
            f.write(f"ID: {r['id']}\nTítulo: {r['titulo']}\nDescrição: {r['descricao']}\nData: {r['data']}\n")
            f.write("-" * 40 + "\n")

    print("📁 Relatórios salvos em 'relatorios.txt'.")

def enviar_email(destinatario, assunto, corpo):
    """Envia os relatórios por e-mail."""
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587
    smtp_user = 'seu email'
    smtp_password = 'sua senha'  # Substitua por variável de ambiente para segurança

    msg = EmailMessage()
    msg['Subject'] = assunto
    msg['From'] = smtp_user
    msg['To'] = destinatario
    msg.set_content(corpo)

    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            print("📧 Email enviado com sucesso!")
    except Exception as e:
        print(f"❌ Erro ao enviar e-mail: {e}")

def enviar_relatorios_por_email():
    """Gera um resumo dos relatórios e envia por e-mail."""
    if not relatorios:
        print("⚠️ Nenhum relatório para enviar.")
        return

    destinatario = input("Digite o e-mail de destino: ")
    assunto = "Relatórios de Ronda"
    corpo = "\n\n".join(
        [f"ID: {r['id']}\nTítulo: {r['titulo']}\nDescrição: {r['descricao']}\nData: {r['data']}" for r in relatorios]
    )

    enviar_email(destinatario, assunto, corpo)

# ---------- Menu Principal ----------

def menu():
    carregar_json()
    while True:
        print("\n===== Sistema de Cadastro de Relatórios =====")
        print("1. Cadastrar Relatório")
        print("2. Listar Relatórios")
        print("3. Buscar Relatório por ID")
        print("4. Buscar Relatório por Título")
        print("5. Exportar Relatórios para DOCX")
        print("6. Exportar Relatórios para TXT")
        print("7. Enviar Relatórios por Email")
        print("8. Sair")

        opcao = input("Escolha uma opção: ")

        if opcao == '1':
            cadastrar_relatorio()
        elif opcao == '2':
            listar_relatorios()
        elif opcao == '3':
            buscar_relatorio_por_id()
        elif opcao == '4':
            buscar_por_titulo()
        elif opcao == '5':
            salvar_em_docx()
        elif opcao == '6':
            salvar_em_txt()
        elif opcao == '7':
            enviar_relatorios_por_email()
        elif opcao == '8':
            print("👋 Saindo do sistema...")
            break
        else:
            print("❌ Opção inválida. Tente novamente.")

# ---------- Execução ----------

if __name__ == "__main__":
    menu()
